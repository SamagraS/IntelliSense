from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import logging
import re

logger = logging.getLogger(__name__)

# ============================================================================
# HELPER FUNCTIONS FOR VISUAL ENHANCEMENT
# ============================================================================

def set_cell_background(cell, color_hex):
    """Set background color of a table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)

def get_decision_color(decision_band):
    """Return color based on decision band"""
    decision_band = str(decision_band).upper()
    if "APPROVE" in decision_band or "SANCTION" in decision_band:
        return 'D4EDDA'  # Light green
    elif "REFER" in decision_band or "REVIEW" in decision_band:
        return 'FFF3CD'  # Light yellow/amber
    else:  # REJECT or DECLINE
        return 'F8D7DA'  # Light red

def get_score_color(score, threshold_low=5.0, threshold_high=7.0):
    """Return color based on score"""
    try:
        score = float(score)
        if score >= threshold_high:
            return RGBColor(40, 167, 69)  # Green
        elif score >= threshold_low:
            return RGBColor(255, 193, 7)  # Amber
        else:
            return RGBColor(220, 53, 69)  # Red
    except:
        return RGBColor(0, 0, 0)  # Black default

def add_score_bar(paragraph, score, max_score=10):
    """Add a visual progress bar with score"""
    try:
        score_val = float(score)
        
        # Add score text
        run = paragraph.add_run(f"{score_val:.1f}/10  ")
        run.bold = True
        run.font.size = Pt(10)
        
        # Add visual bar
        filled = int((score_val / max_score) * 10)
        bar = '█' * filled + '░' * (10 - filled)
        run = paragraph.add_run(bar)
        run.font.size = Pt(10)
        
        # Color based on score
        run.font.color.rgb = get_score_color(score_val)
    except:
        paragraph.add_run(str(score))

def format_header_row(row, bg_color='2C3E50'):
    """Format a table header row with dark background and white text"""
    for cell in row.cells:
        set_cell_background(cell, bg_color)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
                run.font.size = Pt(11)

def add_alert_box(doc, title, message, alert_type='warning'):
    """Add a colored alert box for critical information"""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    
    cell = table.rows[0].cells[0]
    
    if alert_type == 'critical':
        set_cell_background(cell, 'F8D7DA')  # Light red
        icon = '⚠ CRITICAL: '
        title_color = RGBColor(114, 28, 36)
    elif alert_type == 'warning':
        set_cell_background(cell, 'FFF3CD')  # Light yellow
        icon = '⚠ WARNING: '
        title_color = RGBColor(133, 100, 4)
    else:
        set_cell_background(cell, 'D1ECF1')  # Light blue
        icon = 'ℹ INFO: '
        title_color = RGBColor(12, 84, 96)
    
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(icon + title)
    run.bold = True
    run.font.color.rgb = title_color
    run.font.size = Pt(11)
    
    paragraph.add_run('\n' + message)
    
    # Add some padding
    cell.paragraphs[0].paragraph_format.space_before = Pt(6)
    cell.paragraphs[0].paragraph_format.space_after = Pt(6)

def clean_text(text: str):
    """Clean markdown formatting from text"""
    text = re.sub(r"#{1,6}\s*", "", text)
    text = text.replace("---", "")
    text = text.replace("*", "")
    return text.strip()

# ============================================================================
# TABLE BUILDERS
# ============================================================================

def add_key_value_table(doc, title, data: dict, highlight_row=None):
    """Add a key-value table with optional row highlighting"""
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    
    # Format header
    format_header_row(table.rows[0])
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Value"
    
    for k, v in data.items():
        row = table.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)
        
        # Highlight specific rows
        if highlight_row and k == highlight_row:
            if k == "Decision":
                set_cell_background(row[1], get_decision_color(v))
            elif k == "Composite Score":
                try:
                    score = float(v)
                    if score >= 7.0:
                        set_cell_background(row[1], 'D4EDDA')
                    elif score >= 5.5:
                        set_cell_background(row[1], 'FFF3CD')
                    else:
                        set_cell_background(row[1], 'F8D7DA')
                except:
                    pass

def add_financial_table(doc, financials: dict):
    """Add financial snapshot table with metric highlighting"""
    doc.add_heading("Financial Snapshot", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    
    # Format header
    format_header_row(table.rows[0])
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    
    metrics = [
        ("Revenue FY2024", financials.get("revenue_cr", {}).get("FY2024")),
        ("EBITDA FY2024", financials.get("ebitda_cr", {}).get("FY2024")),
        ("Net Profit FY2024", financials.get("net_profit_cr", {}).get("FY2024")),
        ("DSCR", financials.get("dscr")),
        ("Debt / Equity", financials.get("debt_to_equity")),
        ("Current Ratio", financials.get("current_ratio")),
    ]
    
    for name, val in metrics:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = str(val) if val is not None else "N/A"
        
        # Highlight critical metrics below threshold
        if name == "DSCR" and val is not None:
            try:
                dscr_val = float(val)
                if dscr_val < 1.25:
                    set_cell_background(row[1], 'F8D7DA')  # Red
                    run = row[1].paragraphs[0].runs[0]
                    run.bold = True
                    run.font.color.rgb = RGBColor(114, 28, 36)
            except:
                pass
        
        if name == "Current Ratio" and val is not None:
            try:
                cr_val = float(val)
                if cr_val < 1.0:
                    set_cell_background(row[1], 'FFF3CD')  # Yellow
            except:
                pass

def add_five_cs_table(doc, five_cs: dict):
    """Add Five Cs scorecard with visual score bars"""
    doc.add_heading("Five Cs Scorecard", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    
    # Format header
    format_header_row(table.rows[0])
    hdr = table.rows[0].cells
    hdr[0].text = "Credit Dimension"
    hdr[1].text = "Score"
    hdr[2].text = "Weight"
    
    for c, val in five_cs.items():
        row = table.add_row().cells
        row[0].text = c
        
        # Add visual score bar
        score = val.get("score")
        if score is not None:
            add_score_bar(row[1].paragraphs[0], score)
        else:
            row[1].text = "N/A"
        
        row[2].text = str(val.get("c_level_weight"))

def add_swot_matrix(doc, swot_text: str):
    """Add enhanced SWOT matrix with colored quadrants"""
    doc.add_heading("SWOT Matrix", level=2)
    
    strengths = []
    weaknesses = []
    opportunities = []
    threats = []
    
    lines = swot_text.splitlines()
    for l in lines:
        l_lower = l.lower()
        if "strength" in l_lower:
            strengths.append(l)
        if "weakness" in l_lower:
            weaknesses.append(l)
        if "opportunit" in l_lower:
            opportunities.append(l)
        if "threat" in l_lower:
            threats.append(l)
    
    # Create 2x2 SWOT matrix
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    
    # Strengths
    cell_s = table.rows[0].cells[0]
    set_cell_background(cell_s, 'D1F2EB')  # Light green
    p = cell_s.paragraphs[0]
    run = p.add_run("Strengths")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(21, 87, 36)
    cell_s.add_paragraph("\n".join(strengths) if strengths else "N/A")
    
    # Weaknesses
    cell_w = table.rows[0].cells[1]
    set_cell_background(cell_w, 'FCF3CF')  # Light yellow
    p = cell_w.paragraphs[0]
    run = p.add_run("Weaknesses")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(133, 100, 4)
    cell_w.add_paragraph("\n".join(weaknesses) if weaknesses else "N/A")
    
    # Opportunities
    cell_o = table.rows[1].cells[0]
    set_cell_background(cell_o, 'D6EAF8')  # Light blue
    p = cell_o.paragraphs[0]
    run = p.add_run("Opportunities")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(12, 84, 96)
    cell_o.add_paragraph("\n".join(opportunities) if opportunities else "N/A")
    
    # Threats
    cell_t = table.rows[1].cells[1]
    set_cell_background(cell_t, 'FADBD8')  # Light red
    p = cell_t.paragraphs[0]
    run = p.add_run("Threats")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(114, 28, 36)
    cell_t.add_paragraph("\n".join(threats) if threats else "N/A")

def detect_and_add_alerts(doc, case_json):
    """Detect critical issues and add alert boxes"""
    alerts = []
    
    # Check DSCR
    financials = case_json.get("financial_summary", {})
    dscr = financials.get("dscr")
    if dscr is not None:
        try:
            if float(dscr) < 1.25:
                alerts.append(("DSCR Below Threshold", 
                              f"DSCR of {dscr} is below the policy minimum of 1.25", 
                              "critical"))
        except:
            pass
    
    # Check for auditor issues (from sections if available)
    five_cs = case_json.get("five_cs_scores", {})
    character_score = five_cs.get("Character", {}).get("score")
    if character_score is not None:
        try:
            if float(character_score) < 5.0:
                alerts.append(("Governance Concerns", 
                              f"Character score of {character_score} indicates significant governance issues", 
                              "warning"))
        except:
            pass
    
    # Add alerts if any
    for title, msg, alert_type in alerts:
        add_alert_box(doc, title, msg, alert_type)
        doc.add_paragraph()  # Spacing

# ============================================================================
# MAIN DOCUMENT BUILDER
# ============================================================================

def build_docx(sections: dict, case_json: dict, output_path):
    """Build the complete CAM document with visual enhancements"""
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading("Credit Appraisal Memo", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Extract metadata
        meta = case_json.get("case_metadata", {})
        decision = case_json.get("final_decision", {})
        
        # Decision Snapshot Table
        dashboard = {
            "Company": meta.get("company_name"),
            "Sector": meta.get("sector"),
            "Requested Loan": decision.get("requested_amount_inr"),
            "Suggested Limit": decision.get("suggested_limit_inr"),
            "Composite Score": decision.get("composite_score"),
            "Decision": decision.get("decision_band"),
            "Risk Premium (bps)": decision.get("risk_premium_bps"),
        }
        add_key_value_table(doc, "Decision Snapshot", dashboard, highlight_row="Decision")
        
        # Financial Snapshot
        if "financial_summary" in case_json:
            add_financial_table(doc, case_json["financial_summary"])
        
        # Five Cs Scorecard
        if "five_cs_scores" in case_json:
            add_five_cs_table(doc, case_json["five_cs_scores"])
        
        # Add critical alerts
        detect_and_add_alerts(doc, case_json)
        
        # Define correct section order
        section_order = [
            "Executive Summary",
            "Company and Promoter Profile",
            "SWOT Analysis",  # Moved here
            "Industry and Conditions",
            "Financial Analysis",
            "Five Cs Narrative",
            "Pre-Cognitive Risk Analysis",
            "Proposed Facility and Structure",
            "Recommendation and Rationale",
            "Audit Trail Summary"
        ]
        
        # Add sections in correct order
        for section_name in section_order:
            content = sections.get(section_name)
            if not content:
                continue
            
            # Add section heading
            doc.add_heading(section_name, level=1)
            
            # Clean and add content
            cleaned = clean_text(content)
            paragraphs = re.split(r"\n\s*\n", cleaned)
            
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                doc.add_paragraph(para)
            
            # Add SWOT matrix after SWOT Analysis narrative
            if section_name == "SWOT Analysis":
                add_swot_matrix(doc, cleaned)
        
        # Save document
        doc.save(output_path)
        logger.info(f"Enhanced DOCX saved: {output_path}")
        
    except Exception:
        logger.exception("DOCX build failed")
        raise